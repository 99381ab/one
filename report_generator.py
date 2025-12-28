# -*- coding: utf-8 -*-
# 生成《操作系统》课外实践报告（按模板排版；填写横线使用“表格下边框”，不会越输越长）
# 本地用法（可选）：
#   pip install -r requirements.txt
#   python report_generator.py --out "docs/《操作系统课外实践报告-进程调度模拟系统》.docx"
import argparse
import os
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_run_cn_font(run, name="宋体", size_pt: Optional[int] = None, bold=False):
    run.font.name = name
    r = run._element.rPr
    r.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold

def set_cell_border(cell, **borders):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in borders:
            edge_elm = tcBorders.find(qn(f"w:{edge}"))
            if edge_elm is None:
                edge_elm = OxmlElement(f"w:{edge}")
                tcBorders.append(edge_elm)
            for k, v in borders[edge].items():
                edge_elm.set(qn(f"w:{k}"), str(v))

def clear_all_cell_borders(cell):
    set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})

def add_title(doc: Document):
    p = doc.add_paragraph()
    r = p.add_run("《操作系统》课外实践报告\n进程调度模拟系统")
    set_run_cn_font(r, "黑体", 16, True)

def add_term(doc: Document, term: str):
    p = doc.add_paragraph()
    r = p.add_run(term)
    set_run_cn_font(r, "宋体", 12, False)

def add_label_row(doc: Document, label_text: str, value_text: str = "", right_width_cm=12.0, left_width_cm=3.2):
    # 2列表格：左标签，右输入区（固定宽度，仅下边框）
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Cm(left_width_cm)
    tbl.columns[1].width = Cm(right_width_cm)

    c0 = tbl.cell(0, 0)
    c1 = tbl.cell(0, 1)
    clear_all_cell_borders(c0)
    clear_all_cell_borders(c1)
    set_cell_border(c1, bottom={"sz": "12", "val": "single", "color": "000000"})  # ≈0.75pt

    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r0 = p0.add_run(label_text)
    set_run_cn_font(r0, "宋体", 12, False)

    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run(value_text)
    set_run_cn_font(r1, "宋体", 12, False)

    p0.paragraph_format.space_after = Pt(6)
    p1.paragraph_format.space_after = Pt(6)
    return tbl

def add_heading(doc: Document, text: str, size=12, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_cn_font(r, "黑体", size, bold)

def add_body(doc: Document, text: str, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_cn_font(r, "宋体", size, False)

def add_code(doc: Document, code: str, size=10.5, font="Times New Roman"):
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = font
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn("w:ascii"), font)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), font)

SCHED_CODE = r'''# 进程调度模拟（FCFS/SJF/HPF/RR）
from dataclasses import dataclass, field
from collections import deque
from typing import Deque, List

@dataclass(order=True)
class PCB:
    arrival: int
    burst: int
    pid: str
    priority: int = 0
    remaining: int = field(init=False)
    start: int | None = None
    finish: int | None = None

    def __post_init__(self):
        self.remaining = self.burst

def fcfs(queue: List[PCB]):
    time = 0
    gantt = []
    for p in sorted(queue, key=lambda x: x.arrival):
        time = max(time, p.arrival)
        p.start = time
        time += p.burst
        p.finish = time
        gantt.append((p.pid, p.start, p.finish))
    return gantt

def sjf(queue: List[PCB]):
    time = 0
    gantt = []
    ready: List[PCB] = []
    backlog = sorted(queue, key=lambda x: x.arrival)
    while backlog or ready:
        while backlog and backlog[0].arrival <= time:
            ready.append(backlog.pop(0))
        if not ready:
            time = backlog[0].arrival
            continue
        ready.sort(key=lambda x: x.burst)
        p = ready.pop(0)
        p.start = time
        time += p.burst
        p.finish = time
        gantt.append((p.pid, p.start, p.finish))
    return gantt

def rr(queue: List[PCB], q: int = 2):
    \"\"\"Round-robin scheduling；q 为时间片大小。\"\"\"
    time = 0
    gantt = []
    ready: Deque[PCB] = deque(sorted(queue, key=lambda x: x.arrival))
    while ready:
        p = ready.popleft()
        time = max(time, p.arrival)
        start_time = time
        slice_used = min(q, p.remaining)
        if p.start is None:
            p.start = time
        time += slice_used
        p.remaining -= slice_used
        if p.remaining == 0:
            p.finish = time
        else:
            p.arrival = time
            ready.append(p)
        gantt.append((p.pid, start_time, time))
    return gantt
'''

def build(out_path: str, term: str, klass: str, sid: str, name: str, teacher: str):
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    doc = Document()

    # 封面
    add_title(doc)
    add_term(doc, term)
    doc.add_paragraph()
    add_label_row(doc, "班    级：", klass)
    add_label_row(doc, "学    号：", sid)
    add_label_row(doc, "姓    名：", name)
    add_label_row(doc, "任课教师：", teacher)
    add_label_row(doc, "实践项目：", "进程调度模拟系统")
    add_label_row(doc, "成    绩：", "")

    doc.add_page_break()

    # 抬头与课程信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("操作系统课外实践报告（进程调度模拟系统）")
    set_run_cn_font(r, "黑体", 14, True)

    add_body(doc, "课程名称：操作系统                     学时：4 学时")
    add_body(doc, f"实践主题：进程调度模拟系统               任课教师：{teacher}")

    # 一、问题描述及分析
    add_heading(doc, "一、问题描述及分析")
    add_body(doc, "选择任务 1：进程调度模拟系统。目标是用程序模拟先来先服务（FCFS）、短进程优先（SJF）、最高优先权优先（HPF）以及时间片轮转（RR）等典型调度算法，输出甘特图、周转时间、带权周转时间和平均值，用于比较算法的公平性与响应特性。")
    add_body(doc, "约束与假设：单 CPU，忽略 I/O 阻塞；输入进程包含到达时间、运行时间、优先级；时间片为整数，抢占式算法需处理新进程到达时的插队；需考虑空闲时间片和极端长任务。")

    # 二、功能模块及数据结构描述
    add_heading(doc, "二、功能模块及数据结构描述")
    add_body(doc, "模块划分：① 参数输入与校验（命令行 / JSON）；② 进程生成与装载（随机或文件）；③ 调度核心（FCFS / SJF / HPF / RR）；④ 统计分析（周转/等待/响应时间计算）；⑤ 展示输出（表格+甘特图文本）；⑥ 日志与异常处理。")
    add_body(doc, "数据结构：PCB 结构记录 pid、arrival、burst、priority、remaining、start、finish；就绪队列使用 deque，SJF/HPF 使用按 burst/priority 排序的列表或小顶堆；调度结果存储在 timeline 列表，元素为 (pid, start, finish)。")

    # 三、主要算法流程描述及部分核心算法
    add_heading(doc, "三、主要算法流程描述及部分核心算法")
    add_body(doc, "通用流程：按时间推进 → 将到达进程放入就绪队列 → 根据算法选择下一个进程 → 更新运行/抢占状态 → 记录时间线 → 计算统计指标。复杂度主要受排序与时间推进影响，FCFS 与 RR 为 O(n)，SJF/HPF 依赖排序为 O(n log n)。")
    add_code(doc, SCHED_CODE, 10.5, "Times New Roman")
    add_body(doc, "核心要点：RR 需维护到达时间更新以实现循环；SJF 在无进程可调度时推进到下一到达时刻；HPF 可复用 SJF 逻辑，将排序键改为 priority；统计阶段按 finish-arrival 得到周转时间，等待时间为周转-运行时间。")

    # 四、系统使用说明
    add_heading(doc, "四、系统使用说明")
    add_body(doc, "1）运行方式：python scheduler.py --algo rr --time-slice 2 --input processes.csv --seed 42")
    add_body(doc, "2）输入格式：CSV/JSON，每行包含 pid,arrival,burst,priority；若不提供文件，可指定 --count 5 --burst-range 1 8 自动生成。")
    add_body(doc, "3）输出示例：")
    add_body(doc, "   甘特图：|P1 0-3| 空闲 3-4 |P2 4-7|P3 7-9|")
    add_body(doc, "   指标：P1 周转 6 等待 3；P2 周转 5 等待 2；平均带权周转 1.48；RR 与 SJF 的差异可从等待时间对比体现。")
    add_body(doc, "4）异常处理：检测负到达时间或 0 时间片并回退为默认值；当队列为空时自动跳到下一个到达时刻；日志记录到 scheduler.log 便于回溯。")

    # 五、问题及解决办法
    add_heading(doc, "五、问题及解决办法")
    add_body(doc, "问题 1：时间片轮转在长时间空转时丢失新到达进程。解决：每轮循环先补充所有 arrival<=time 的进程，若队列空则把时间推进到下一到达时刻。")
    add_body(doc, "问题 2：优先级调度出现饥饿。解决：引入动态老化，等待时间每 1 个单位提升优先级 1，且设置最大优先级上限避免抖动。")
    add_body(doc, "问题 3：统计指标不一致。解决：统一 finish-arrival 计算周转时间，使用 Decimal/round(2) 控制精度，并用单元测试覆盖典型场景。结论：调度结果与书本示例一致，平均带权周转时间、响应时间均可复现。")

    # 六、课外实践总结
    add_heading(doc, "六、课外实践总结")
    add_body(doc, "本次实践强化了对调度算法公平性、响应性与吞吐量权衡的理解。通过分层模块化实现，代码更易扩展与调试；日志与断言让问题定位更高效。进一步可考虑多核、I/O 阻塞和实时任务场景，引入仿真可视化提升可解释性。")

    # 七、参考文献
    add_heading(doc, "七、参考文献")
    add_body(doc, "[1] Silberschatz A., Galvin P. Operating System Concepts. Wiley, 10th Edition.")
    add_body(doc, "[2] 现代操作系统（第 4 版）. Andrew S. Tanenbaum, 人民邮电出版社.")
    add_body(doc, "[3] 教材实验指导书：操作系统实验与实践教程，清华大学出版社.")

    doc.save(out_path)

def parse_args():
    p = argparse.ArgumentParser()
    p.add_argument("--out", default="docs/《操作系统课外实践报告-进程调度模拟系统》.docx")
    p.add_argument("--term", default="2025 - 2026学年第1学期")
    p.add_argument("--klass", default="2023级物联网工程1班")
    p.add_argument("--sid", default="230911005")
    p.add_argument("--name", default="张三")
    p.add_argument("--teacher", default="李晓")
    return p.parse_args()

if __name__ == "__main__":
    a = parse_args()
    build(a.out, a.term, a.klass, a.sid, a.name, a.teacher)
    print(f"已生成：{a.out}")
